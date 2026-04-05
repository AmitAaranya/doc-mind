"""MS Word processor — extracts text and tables from .docx / .doc files.

Uses python-docx to walk the document body in XML order, preserving the
reading sequence of paragraphs and tables.  Output is an ``OrderedPDFDocument``
so it plugs directly into the existing chunker pipeline.

Public API
----------
processor = DocxProcessor("report.docx")
doc = processor.extract_ordered()   # -> OrderedPDFDocument
"""

from __future__ import annotations

from pathlib import Path

from app.utils.pdf_processor import (
    BlockType,
    OrderedPageData,
    OrderedPDFDocument,
    TableBlock,
    TableData,
    TextBlock,
)

# Maximum characters accumulated before a virtual page break is inserted.
_PAGE_SIZE = 4_000


class DocxProcessor:
    """Extract text and tables from a .docx or .doc file.

    Parameters
    ----------
    file_path:
        Path to the Word document.

    Raises
    ------
    FileNotFoundError
        If the file does not exist.
    ValueError
        If the file is not a .docx or .doc file.
    """

    def __init__(self, file_path: str | Path) -> None:
        self._path = Path(file_path)
        if not self._path.exists():
            raise FileNotFoundError(f"File not found: {self._path}")
        if self._path.suffix.lower() not in {".docx", ".doc"}:
            raise ValueError(f"Expected a .docx or .doc file, got: {self._path.suffix}")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract_ordered(self) -> OrderedPDFDocument:
        """Return an ``OrderedPDFDocument`` with text and table blocks in body
        reading order.

        Pages contain :class:`TextBlock` and :class:`TableBlock` objects only
        (no :class:`ImageBlock`).  Virtual page breaks are inserted every
        ``_PAGE_SIZE`` characters of accumulated paragraph text.
        """
        import docx  # python-docx; lazy import so the rest of the app never breaks

        document = docx.Document(str(self._path))
        body = document.element.body

        ordered_pages: list[OrderedPageData] = []
        page_number = 1
        current_blocks: list[TextBlock | TableBlock] = []
        current_text_parts: list[str] = []
        current_text_len = 0
        para_idx = 0
        table_idx = 0

        # ------------------------------------------------------------------
        # Inner helpers (closures over the mutable state above)
        # ------------------------------------------------------------------

        def _flush_text() -> None:
            nonlocal current_text_len
            if not current_text_parts:
                return
            merged = "\n".join(current_text_parts).strip()
            if merged:
                current_blocks.append(
                    TextBlock(
                        block_type=BlockType.TEXT,
                        page_number=page_number,
                        order=len(current_blocks),
                        bbox=(0.0, 0.0, 0.0, 0.0),
                        content=merged,
                    )
                )
            current_text_parts.clear()
            current_text_len = 0

        def _commit_page() -> None:
            nonlocal page_number, current_blocks
            _flush_text()
            if not current_blocks:
                return
            for i, blk in enumerate(current_blocks):
                blk.order = i
            ordered_pages.append(
                OrderedPageData(page_number=page_number, blocks=list(current_blocks))
            )
            page_number += 1
            current_blocks.clear()

        # ------------------------------------------------------------------
        # Walk body elements in XML order
        # ------------------------------------------------------------------
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para = (
                    document.paragraphs[para_idx] if para_idx < len(document.paragraphs) else None
                )
                para_idx += 1
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                # Insert a virtual page break when accumulated text is large enough
                if current_text_len + len(text) > _PAGE_SIZE and current_text_parts:
                    _flush_text()
                    _commit_page()
                current_text_parts.append(text)
                current_text_len += len(text)

            elif tag == "tbl":
                tbl = document.tables[table_idx] if table_idx < len(document.tables) else None
                table_idx += 1
                if tbl is None:
                    continue
                rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
                if not rows:
                    continue
                _flush_text()  # commit pending text before the table
                headers = rows[0]
                data_rows = rows[1:]
                current_blocks.append(
                    TableBlock(
                        block_type=BlockType.TABLE,
                        page_number=page_number,
                        order=len(current_blocks),
                        bbox=(0.0, 0.0, 0.0, 0.0),
                        table=TableData(
                            index=table_idx - 1,
                            headers=headers,
                            rows=data_rows,
                        ),
                    )
                )

        # Flush anything left over
        _flush_text()
        if current_blocks:
            for i, blk in enumerate(current_blocks):
                blk.order = i
            ordered_pages.append(
                OrderedPageData(page_number=page_number, blocks=list(current_blocks))
            )

        metadata = {
            "author": document.core_properties.author or "",
            "title": document.core_properties.title or "",
            "format": "docx",
        }

        return OrderedPDFDocument(
            file_path=str(self._path),
            total_pages=len(ordered_pages),
            metadata=metadata,
            pages=ordered_pages,
        )
